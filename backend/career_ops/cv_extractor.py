import io
import json
import re

from openai import AsyncOpenAI

from backend.database import settings

_MODEL = "gpt-4o-mini"

_SYSTEM = """You are an expert HR analyst. Extract career profile information from the provided CV/resume.
Return ONLY valid JSON with no markdown or explanation."""

_PROMPT = """Extract the following fields from this CV. Use empty string "" for text fields not found, empty array [] for list fields not found.

{
  "full_name": "",
  "email": "",
  "phone": "",
  "location": "",
  "linkedin": "",
  "portfolio_url": "",
  "github": "",
  "primary_roles": [],
  "headline": "",
  "exit_story": "",
  "superpowers": [],
  "target_range": "",
  "currency": "USD",
  "minimum_salary": "",
  "location_flexibility": "",
  "country": "",
  "city": "",
  "timezone": "",
  "visa_status": "",
  "deal_breakers": [],
  "must_haves": []
}

Field guidance:
- primary_roles: 2-4 job titles the person is targeting, inferred from recent experience and seniority
- headline: a concise one-line professional summary based on their experience
- exit_story: their unique value proposition based on achievements and skills
- superpowers: 3-6 core technical or strategic strengths
- location: full location string; also split into country and city
- target_range / minimum_salary: only if explicitly mentioned
- location_flexibility: remote/hybrid preference if mentioned
- deal_breakers / must_haves: only if explicitly stated
- linkedin / github / portfolio_url: extract URLs if present"""


def _client() -> AsyncOpenAI:
    return AsyncOpenAI(api_key=settings.openai_api_key.strip())


async def extract_pdf_text_via_openai(content: bytes, filename: str) -> dict:
    """Upload PDF to OpenAI Files API — no local PDF parsing needed."""
    client = _client()

    file_obj = await client.files.create(
        file=(filename, io.BytesIO(content), "application/pdf"),
        purpose="user_data",
    )
    try:
        response = await client.chat.completions.create(
            model=_MODEL,
            temperature=0,
            max_tokens=1200,
            messages=[
                {"role": "system", "content": _SYSTEM},
                {
                    "role": "user",
                    "content": [
                        {"type": "file", "file": {"file_id": file_obj.id}},
                        {"type": "text", "text": _PROMPT},
                    ],
                },
            ],
        )
    finally:
        try:
            await client.files.delete(file_obj.id)
        except Exception:
            pass

    return _parse_response(response.choices[0].message.content)


async def extract_profile_from_cv_text(text: str) -> dict:
    """Send plain text to OpenAI for DOCX / TXT files."""
    client = _client()

    response = await client.chat.completions.create(
        model=_MODEL,
        temperature=0,
        max_tokens=1200,
        messages=[
            {"role": "system", "content": _SYSTEM},
            {"role": "user", "content": f"{_PROMPT}\n\nCV TEXT:\n{text[:8000]}"},
        ],
    )
    return _parse_response(response.choices[0].message.content)


def extract_docx_text(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    return "\n".join(paragraphs)


def _parse_response(raw: str) -> dict:
    raw = raw.strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


def count_filled(profile: dict) -> int:
    count = 0
    for v in profile.values():
        if isinstance(v, str) and v.strip():
            count += 1
        elif isinstance(v, list) and v:
            count += 1
    return count
